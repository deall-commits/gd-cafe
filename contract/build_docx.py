#!/usr/bin/env python3
"""계약서 마크다운 원본을 서명 가능한 Word(.docx) 문서로 변환한다.

사용법:  python3 contract/build_docx.py
같은 폴더의 *.md 를 모두 읽어 동일한 이름의 .docx 를 만든다.

마크다운 규약
  # 제목                 문서 제목 (가운데 정렬)
  [부제] 계약기간 1년     제목 아래 부제 (가운데 정렬)
  > 문단                 전문(前文) 문단
  [쪽나눔]               페이지 나누기
  ## 제1조 (목적)        조 제목
  ①/1./평문             조문 본문
  ---                    구분선 (서명란 시작)
  [표] 라벨 | 라벨       2열 표 (첫 행이 머리글)
  [표=] 항목 | 내용      2열 표 (왼쪽 칸이 항목명)
  [표-] 칸 | 칸          2열 표 (강조 없음)
  [양식표] 머리글,머리글 | 12   빈 칸이 12행인 기재용 표
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = "맑은 고딕"
TITLE_FONT = "맑은 고딕"


def _font(run, name=BODY_FONT, size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)
    return run


def _setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)


def _para(doc, text, *, align=None, size=10.5, bold=False, before=0, after=4,
          indent=0.0, hanging=0.0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.5
    if indent:
        pf.left_indent = Cm(indent)
    if hanging:
        pf.first_line_indent = Cm(-hanging)
    if text:
        _font(p.add_run(text), size=size, bold=bold)
    return p


def _rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    pbdr = p._p.get_or_add_pPr().makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"), {})
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)


def _sign_table(doc, rows, emphasis="head"):
    """rows: [(왼쪽 칸, 오른쪽 칸), ...]

    emphasis "head" 첫 행을 굵게, "label" 왼쪽 칸을 굵게, "none" 강조 없음.
    """
    widths = (Cm(4.2), Cm(11.4)) if emphasis == "label" else (Cm(7.8), Cm(7.8))
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r, (left, right) in enumerate(rows):
        for c, text in enumerate((left, right)):
            cell = table.cell(r, c)
            cell.width = widths[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.4
            bold = (emphasis == "head" and r == 0) or (emphasis == "label" and c == 0)
            _font(p.add_run(text), size=10.5, bold=bold)
    return table


def _form_table(doc, headers, rows):
    table = doc.add_table(rows=rows + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, head in enumerate(headers):
        p = table.cell(0, c).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(head), size=10, bold=True)
    for r in range(1, rows + 1):
        for c in range(len(headers)):
            p = table.cell(r, c).paragraphs[0]
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            _font(p.add_run(""), size=10)
    return table


TABLE_RE = re.compile(r"^\[표([=-]?)\]")
TABLE_EMPHASIS = {"": "head", "=": "label", "-": "none"}

ARTICLE_RE = re.compile(r"^##\s+(제\s*\d+\s*조.*)$")
ITEM_RE = re.compile(r"^[①-⑳]")
SUBITEM_RE = re.compile(r"^\d+\.\s")
DASH_RE = re.compile(r"^-\s")


def build(md_path: Path, docx_path: Path):
    doc = Document()
    _setup(doc)

    lines = md_path.read_text(encoding="utf-8").split("\n")
    i = 0
    in_sign = False
    sign_rows = []

    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        if not line.strip():
            continue

        if line.startswith("<!--"):
            continue

        if line.startswith("# "):
            _para(doc, line[2:].strip(), align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=19, bold=True, before=0, after=4)
            continue

        if line.startswith("[부제]"):
            _para(doc, line[4:].strip(), align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, bold=False, before=0, after=20)
            continue

        if line.strip() == "[쪽나눔]":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            continue

        if line.strip() == "---":
            _rule(doc)
            in_sign = True
            continue

        if line.startswith("### "):
            _para(doc, line[4:].strip(), size=11.5, bold=True, before=14, after=6)
            continue

        m = ARTICLE_RE.match(line)
        if m:
            _para(doc, m.group(1).strip(), size=11.5, bold=True, before=13, after=5)
            continue

        if line.startswith("> "):
            _para(doc, line[2:].strip(), align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  before=2, after=8)
            continue

        if line.startswith("[양식표]"):
            spec, _, count = line[5:].rpartition("|")
            headers = [h.strip() for h in spec.split(",") if h.strip()]
            _form_table(doc, headers, int(count.strip() or 10))
            continue

        m = TABLE_RE.match(line)
        if m:
            marker, emphasis = m.group(0), TABLE_EMPHASIS[m.group(1)]

            def _row(text):
                cells = [c.strip() for c in text[len(marker):].split("|")]
                return (cells[0], cells[1] if len(cells) > 1 else "")

            sign_rows.append(_row(line))
            # 같은 표에 속하는 뒤따르는 줄을 모두 모은다
            while i < len(lines) and lines[i].startswith(marker):
                sign_rows.append(_row(lines[i]))
                i += 1
            _sign_table(doc, sign_rows, emphasis)
            sign_rows = []
            continue

        text = line.strip()
        if ITEM_RE.match(text):
            _para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  indent=0.6, hanging=0.6, after=3)
        elif SUBITEM_RE.match(text):
            _para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  indent=1.2, hanging=0.55, after=2)
        elif DASH_RE.match(text):
            _para(doc, "· " + text[2:], align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  indent=1.2, hanging=0.4, after=2)
        else:
            _para(doc, text,
                  align=WD_ALIGN_PARAGRAPH.CENTER if in_sign else WD_ALIGN_PARAGRAPH.JUSTIFY,
                  after=4)

    doc.save(str(docx_path))
    print(f"  {docx_path.name}")


def main():
    here = Path(__file__).resolve().parent
    targets = sorted(p for p in here.glob("*.md") if not p.name.startswith("README"))
    if not targets:
        print("변환할 .md 파일이 없습니다.", file=sys.stderr)
        return 1
    print("Word 문서 생성:")
    for md in targets:
        build(md, md.with_suffix(".docx"))
    return 0


if __name__ == "__main__":
    sys.exit(main())
